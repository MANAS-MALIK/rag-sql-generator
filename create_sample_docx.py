from docx import Document

doc = Document()

# Add sample queries for Kate
doc.add_paragraph("Client: Kate | Query: SKU Status")
doc.add_paragraph("SELECT sku, status, CASE WHEN last_sold < DATE_SUB(NOW(), INTERVAL 3 YEAR) THEN 'OLD' ELSE 'ACTIVE' END AS tag FROM inventory WHERE client = 'Kate';")

doc.add_paragraph("")

doc.add_paragraph("Client: Kate | Query: Revenue Summary")
doc.add_paragraph("SELECT DATE_FORMAT(order_date, '%Y-%m') AS month, SUM(amount) AS revenue FROM orders WHERE client = 'Kate' GROUP BY month ORDER BY month DESC;")

doc.add_paragraph("")

doc.add_paragraph("Client: Kate | Query: Top Sellers")
doc.add_paragraph("SELECT product_name, SUM(quantity) AS total_sold FROM sales WHERE client = 'Kate' GROUP BY product_name ORDER BY total_sold DESC LIMIT 10;")

doc.add_paragraph("")

# Add sample queries for Nike
doc.add_paragraph("Client: Nike | Query: SKU Status")
doc.add_paragraph("SELECT product_id, stock_status, CASE WHEN last_order_date < DATE_SUB(CURDATE(), INTERVAL 2 YEAR) THEN 'DISCONTINUED' ELSE 'IN_STOCK' END AS product_tag FROM products WHERE brand = 'Nike';")

doc.add_paragraph("")

doc.add_paragraph("Client: Nike | Query: Revenue Summary")
doc.add_paragraph("SELECT YEAR(sale_date) AS year, MONTH(sale_date) AS month, SUM(total_price) AS monthly_revenue FROM transactions WHERE brand = 'Nike' GROUP BY year, month ORDER BY year DESC, month DESC;")

doc.save('sample_queries.docx')
print("✅ Created sample_queries.docx with 5 sample SQL queries")
